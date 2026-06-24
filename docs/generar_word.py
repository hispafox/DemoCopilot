"""Genera analisis-diseño.docx a partir del contenido del análisis del proyecto."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "analisis-diseño.docx")


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table_from_rows(doc, headers, rows, col_widths_cm):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Cabecera
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = Cm(col_widths_cm[i])
        set_cell_bg(cell, "2E75B6")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    # Filas de datos
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, text in enumerate(row_data):
            cell = row.cells[ci]
            cell.width = Cm(col_widths_cm[ci])
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9)

    return table


def add_code_block(doc, code: str):
    for line in code.strip().splitlines():
        p = doc.add_paragraph(style="No Spacing")
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F4F4F4")
        pPr.append(shd)
    doc.add_paragraph()


def build():
    doc = Document()

    # ----- Estilos globales -----
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    for lvl, (size, bold) in enumerate([(18, True), (14, True), (12, True)], start=1):
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = "Arial"
        h.font.size = Pt(size)
        h.font.bold = bold
        h.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # ===== PORTADA =====
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("Análisis y Diseño")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run("Lista de Tareas — DemoCopilot")
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run("Junio 2026").font.color.rgb = RGBColor(0x90, 0x90, 0x90)

    doc.add_page_break()

    # ===== 1. OBJETIVO =====
    doc.add_heading("1. Objetivo del proyecto", level=1)
    doc.add_paragraph(
        "Aplicación web CRUD de gestión de tareas personales construida como demo didáctica "
        "para el curso de GitHub Copilot. Permite crear, consultar, actualizar y eliminar tareas, "
        "con soporte de plantillas reutilizables y tareas repetitivas con generación automática "
        "de la siguiente ocurrencia al completar. El objetivo prioritario es la claridad del "
        "código sobre la sofisticación arquitectónica."
    )

    # ===== 2. STACK =====
    doc.add_heading("2. Stack tecnológico", level=1)
    add_table_from_rows(
        doc,
        headers=["Tecnología", "Versión", "Rol"],
        rows=[
            ["ASP.NET Core", "8", "API REST / backend (Controllers)"],
            ["Entity Framework Core", "8", "ORM / acceso a datos"],
            ["SQLite", "—", "Base de datos embebida, sin infraestructura externa"],
            ["xUnit + Moq", "—", "Tests unitarios"],
        ],
        col_widths_cm=[5, 3, 9],
    )
    doc.add_paragraph()

    # ===== 3. ARQUITECTURA =====
    doc.add_heading("3. Arquitectura de capas", level=1)
    doc.add_paragraph(
        "Estructura de capas plana y legible en pantalla, sin patrones complejos (sin CQRS ni Mediator)."
    )
    add_table_from_rows(
        doc,
        headers=["Capa", "Carpeta", "Responsabilidad"],
        rows=[
            ["Modelos de dominio", "Models/", "Entidades del negocio (TodoItem, PlantillaTarea, TipoRecurrencia)"],
            ["Acceso a datos", "Data/", "DbContext y configuración de EF Core"],
            ["Lógica de negocio", "LogicaNegocio/", "IXxxLogica + XxxLogica — reglas de negocio y acceso a datos"],
            ["Servicios", "Services/", "IXxxService + XxxService — orquestación y mapeo DTO ↔ entidad"],
            ["API / Controladores", "Controllers/", "Orquestación HTTP, sin lógica de negocio"],
            ["Tests", "Tests/", "Proyecto xUnit separado"],
        ],
        col_widths_cm=[4, 4, 9],
    )
    doc.add_paragraph()

    doc.add_heading("Flujo de llamadas", level=2)
    add_code_block(doc, "Controller → [DTO] → Service → [Entidad] → LogicaNegocio → DbContext")

    doc.add_heading("Reglas de diseño", level=2)
    for rule in [
        "Inyección de dependencias por constructor, nunca new directo de servicios.",
        "async/await en todos los métodos que accedan a base de datos.",
        "Los controladores solo orquestan — sin lógica de negocio dentro de ellos.",
        "Prefijo I para interfaces: ITodoService, IPlantillaService, etc.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(rule).font.size = Pt(11)

    doc.add_paragraph()

    # ===== 4. MODELO DE DATOS =====
    doc.add_heading("4. Modelo de datos", level=1)

    # TodoItem
    doc.add_heading("TodoItem", level=2)
    add_table_from_rows(
        doc,
        headers=["Campo", "Tipo", "Descripción"],
        rows=[
            ["Id", "int", "Clave primaria, autogenerada"],
            ["Title", "string", "Título o descripción de la tarea. Requerido"],
            ["IsCompleted", "bool", "Indica si la tarea está completada. Por defecto false"],
            ["CreatedAt", "DateTime", "Fecha y hora de creación, asignada al crear"],
            ["EsRepetitiva", "bool", "Indica si la tarea se repite automáticamente al completarse"],
            ["Recurrencia", "TipoRecurrencia?", "Periodicidad: Diaria, Semanal o Mensual. null si no es repetitiva"],
            ["ProximaFecha", "DateTime?", "Fecha calculada para la siguiente ocurrencia"],
            ["PlantillaId", "int?", "FK opcional a PlantillaTarea"],
            ["UsuarioAsignadoId", "int?", "FK opcional a UsuarioAsignado"],
        ],
        col_widths_cm=[4.5, 4.5, 8],
    )
    doc.add_paragraph()
    add_code_block(doc, """\
public class TodoItem
{
    public int Id { get; set; }
    public string Title { get; set; } = string.Empty;
    public bool IsCompleted { get; set; }
    public DateTime CreatedAt { get; set; }
    public bool EsRepetitiva { get; set; }
    public TipoRecurrencia? Recurrencia { get; set; }
    public DateTime? ProximaFecha { get; set; }
    public int? PlantillaId { get; set; }
    public PlantillaTarea? Plantilla { get; set; }
    public int? UsuarioAsignadoId { get; set; }
    public UsuarioAsignado? UsuarioAsignado { get; set; }
}""")

    # UsuarioAsignado
    doc.add_heading("UsuarioAsignado", level=2)
    doc.add_paragraph("Usuario que puede ser asignado a una o varias tareas.")
    add_table_from_rows(
        doc,
        headers=["Campo", "Tipo", "Descripción"],
        rows=[
            ["Id", "int", "Clave primaria, autogenerada"],
            ["Nombre", "string", "Nombre completo del usuario. Requerido, máx. 100 caracteres"],
            ["Email", "string", "Dirección de correo electrónico. Requerido, único, máx. 200 caracteres"],
        ],
        col_widths_cm=[4.5, 4.5, 8],
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Relación: un UsuarioAsignado puede tener muchas tareas (TodoItem). "
        "La FK en TodoItem es opcional (SetNull al eliminar el usuario)."
    )

    # PlantillaTarea
    doc.add_heading("PlantillaTarea", level=2)
    doc.add_paragraph("Plantilla reutilizable desde la que se pueden generar nuevas tareas con valores predefinidos.")
    add_table_from_rows(
        doc,
        headers=["Campo", "Tipo", "Descripción"],
        rows=[
            ["Id", "int", "Clave primaria, autogenerada"],
            ["Titulo", "string", "Título por defecto de la tarea generada. Requerido"],
            ["EsRepetitiva", "bool", "Si las tareas generadas desde esta plantilla serán repetitivas"],
            ["Recurrencia", "TipoRecurrencia?", "Periodicidad por defecto para las tareas generadas"],
        ],
        col_widths_cm=[4.5, 4.5, 8],
    )
    doc.add_paragraph()

    # TipoRecurrencia
    doc.add_heading("TipoRecurrencia (enum)", level=2)
    add_code_block(doc, """\
public enum TipoRecurrencia
{
    Diaria,
    Semanal,
    Mensual
}""")

    # ===== 5. ENDPOINTS =====
    doc.add_heading("5. Endpoints API REST", level=1)

    doc.add_heading("Tareas — /api/tareas", level=2)
    add_table_from_rows(
        doc,
        headers=["Verbo", "Ruta", "Descripción", "OK", "Error"],
        rows=[
            ["GET", "/api/tareas", "Listar todas las tareas", "200 + array", "—"],
            ["GET", "/api/tareas/{id}", "Obtener tarea por ID", "200 + objeto", "404"],
            ["POST", "/api/tareas", "Crear nueva tarea", "201 + objeto", "400"],
            ["PUT", "/api/tareas/{id}", "Actualizar título o estado", "200 + objeto", "404 / 400"],
            ["DELETE", "/api/tareas/{id}", "Eliminar tarea", "204", "404"],
            ["POST", "/api/tareas/{id}/completar", "Completar; si es repetitiva genera siguiente ocurrencia", "200 + objeto", "404"],
        ],
        col_widths_cm=[2.5, 5.5, 5.5, 2, 2],
    )
    doc.add_paragraph()

    doc.add_heading("Usuarios asignados — /api/usuariosasignados", level=2)
    add_table_from_rows(
        doc,
        headers=["Verbo", "Ruta", "Descripción", "OK", "Error"],
        rows=[
            ["GET", "/api/usuariosasignados", "Listar todos los usuarios", "200 + array", "—"],
            ["GET", "/api/usuariosasignados/{id}", "Obtener usuario por ID", "200 + objeto", "404"],
            ["POST", "/api/usuariosasignados", "Crear usuario nuevo", "201 + objeto", "400"],
            ["PUT", "/api/usuariosasignados/{id}", "Actualizar nombre o email", "200 + objeto", "404 / 400"],
            ["DELETE", "/api/usuariosasignados/{id}", "Eliminar usuario (tareas quedan sin usuario)", "204", "404"],
        ],
        col_widths_cm=[2.5, 5.5, 5.5, 2, 2],
    )
    doc.add_paragraph()

    doc.add_heading("Plantillas — /api/plantillas", level=2)
    add_table_from_rows(
        doc,
        headers=["Verbo", "Ruta", "Descripción", "OK", "Error"],
        rows=[
            ["GET", "/api/plantillas", "Listar todas las plantillas", "200 + array", "—"],
            ["GET", "/api/plantillas/{id}", "Obtener plantilla por ID", "200 + objeto", "404"],
            ["POST", "/api/plantillas", "Crear nueva plantilla", "201 + objeto", "400"],
            ["PUT", "/api/plantillas/{id}", "Actualizar plantilla", "200 + objeto", "404 / 400"],
            ["DELETE", "/api/plantillas/{id}", "Eliminar plantilla", "204", "404"],
            ["POST", "/api/plantillas/{id}/instanciar", "Crear tarea a partir de la plantilla", "201 + objeto", "404"],
        ],
        col_widths_cm=[2.5, 5.5, 5.5, 2, 2],
    )
    doc.add_paragraph()

    # ===== 6. DATOS DE EJEMPLO =====
    doc.add_heading("6. Datos de ejemplo", level=1)
    doc.add_paragraph(
        "El seeder se invoca desde Program.cs al arrancar la aplicación y solo inserta datos si la tabla está vacía."
    )
    add_table_from_rows(
        doc,
        headers=["Entidad", "Registros de ejemplo"],
        rows=[
            ["UsuarioAsignado", "Ana García (ana@demo.com), Carlos López (carlos@demo.com)"],
            ["PlantillaTarea", '"Revisión semanal" (semanal), "Backup diario" (diaria)'],
            ["TodoItem", '"Configurar entorno" (completada), "Revisar PR pendientes" (repetitiva semanal, asignada a Ana)'],
        ],
        col_widths_cm=[5, 12],
    )
    doc.add_paragraph()

    # ===== 7. DECISIONES =====
    doc.add_heading("7. Decisiones de diseño", level=1)
    decisiones = [
        ("SQLite sobre SQL Server", "Base de datos embebida, sin instalación ni configuración de servidor. Ideal para una demo local y didáctica."),
        ("Controllers sobre Minimal API", "Los controladores con atributos son más explícitos y fáciles de leer en pantalla durante una presentación."),
        ("Inyección de dependencias por constructor", "Patrón estándar de .NET; desacopla la lógica de negocio y facilita los tests con mocks."),
        ("Interfaz ITodoService", "Permite sustituir la implementación por un mock en los tests sin tocar el controlador."),
        ("Recurrencia calculada en el servicio", "La lógica de generar la siguiente ocurrencia vive en TodoService.Completar(), no en el controlador ni en el modelo."),
        ("POST /completar en lugar de PUT", "La acción de completar una tarea repetitiva tiene un efecto secundario (crear nueva ocurrencia), por lo que merece su propio endpoint de acción."),
        ("Plantilla como entidad independiente", "PlantillaTarea es una entidad propia, no un campo de TodoItem, para permitir gestionar y reutilizar plantillas sin acoplarlas al ciclo de vida de las tareas."),
        ("Sin paginación ni autenticación", "Fuera del alcance de la demo; añadir solo si se solicita explícitamente."),
        ("Código en castellano", "Nombres de clases, métodos y variables en español para coherencia con el público hispanohablante del curso."),
    ]
    for titulo, desc in decisiones:
        p = doc.add_paragraph(style="List Bullet")
        run_bold = p.add_run(titulo + ": ")
        run_bold.bold = True
        run_bold.font.size = Pt(11)
        p.add_run(desc).font.size = Pt(11)

    doc.add_paragraph()

    # ===== 8. PENDIENTES =====
    doc.add_heading("8. Pendientes / Preguntas abiertas", level=1)
    pendientes = [
        ("Frontend", "¿Razor Pages o React + Vite? Pendiente de decisión en el curso."),
        ("Paginación", "No incluida. Si el volumen de tareas crece, se añadiría skip/take al endpoint GET."),
        ("Autenticación", "No incluida. Posible extensión futura con JWT o cookies."),
        ("Filtros y búsqueda", "No incluidos. Posible extensión: GET /api/tareas?completada=true."),
        ("Soft delete vs hard delete", "Actualmente se borra físicamente. Podría añadirse campo DeletedAt si se requiere historial."),
        ("Job de recurrencia automática", "Actualmente la siguiente ocurrencia se genera solo al llamar a POST /completar. Si se requiere generación automática, habría que añadir un BackgroundService."),
        ("Recurrencia personalizada", "El enum TipoRecurrencia cubre Diaria/Semanal/Mensual. Para patrones más complejos habría que extender el modelo."),
        ("Herencia de cambios plantilla → tareas", "Si se modifica una plantilla, las tareas ya creadas no se actualizan automáticamente. Comportamiento a definir si se requiere propagación."),
    ]
    for titulo, desc in pendientes:
        p = doc.add_paragraph(style="List Bullet")
        run_bold = p.add_run(titulo + ": ")
        run_bold.bold = True
        run_bold.font.size = Pt(11)
        p.add_run(desc).font.size = Pt(11)

    doc.save(OUTPUT)
    print(f"Documento generado: {OUTPUT}")


if __name__ == "__main__":
    build()
