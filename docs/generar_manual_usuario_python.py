"""
Generador del Manual de Usuario en formato Word (.docx)
Utiliza python-docx para crear un documento profesional con formato.

Instalar: pip install python-docx
Ejecutar: python generar_manual_usuario_python.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def crear_portada(doc):
    """Crea la portada del manual"""
    # Espacio superior
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Título principal
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("Manual de Usuario")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 117, 182)  # #2E75B6
    
    # Subtítulo
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run("Lista de Tareas — DemoCopilot")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    
    # Versión
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("Versión 2.0")
    run.font.size = Pt(14)
    
    # Fecha
    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fecha.add_run("26 de junio de 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Salto de página
    doc.add_page_break()

def agregar_placeholder(doc, texto):
    """Agrega un placeholder visual para imágenes"""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.italic = True
    run.font.color.rgb = RGBColor(136, 136, 136)
    run.font.size = Pt(11)
    # Agregar un espacio después del placeholder
    doc.add_paragraph()

def crear_seccion_intro(doc):
    """Crea la sección de Introducción"""
    doc.add_heading('1. Introducción', level=1)
    
    doc.add_heading('¿Qué es esta aplicación?', level=2)
    doc.add_paragraph(
        'Lista de Tareas es una aplicación web que te permite organizar y gestionar tus tareas '
        'pendientes de forma simple y efectiva. Podrás crear tareas, asignarlas a usuarios, '
        'clasificarlas por categorías, y configurar tareas que se repiten automáticamente.'
    )
    
    doc.add_heading('¿Para quién es?', level=2)
    doc.add_paragraph('Esta aplicación está diseñada para cualquier persona o equipo que necesite:')
    doc.add_paragraph('Llevar un registro de tareas pendientes', style='List Bullet')
    doc.add_paragraph('Organizar el trabajo por categorías (Trabajo, Personal, Urgente, etc.)', style='List Bullet')
    doc.add_paragraph('Asignar tareas a miembros del equipo', style='List Bullet')
    doc.add_paragraph('Automatizar tareas recurrentes (reuniones semanales, backups diarios, etc.)', style='List Bullet')
    
    doc.add_heading('¿Qué puedes hacer con ella?', level=2)
    items = [
        ('Crear y gestionar tareas', 'añade nuevas tareas, edita su contenido, marca como completadas o elimínalas'),
        ('Organizar con categorías', 'agrupa tus tareas por temas con colores visuales'),
        ('Asignar responsables', 'vincula tareas a usuarios específicos'),
        ('Automatizar repeticiones', 'crea tareas que se regeneran automáticamente al completarse (diarias, semanales o mensuales)'),
        ('Usar plantillas', 'guarda configuraciones de tareas frecuentes para crearlas rápidamente'),
    ]
    for titulo, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(titulo).bold = True
        p.add_run(f': {desc}')
    
    doc.add_page_break()

def crear_seccion_primeros_pasos(doc):
    """Crea la sección de Primeros Pasos"""
    doc.add_heading('2. Primeros pasos', level=1)
    
    doc.add_heading('Acceder a la aplicación', level=2)
    doc.add_paragraph('Abre tu navegador web (Chrome, Firefox, Edge, Safari)', style='List Number')
    doc.add_paragraph('Introduce la dirección: http://localhost:5173 (en desarrollo) o la URL proporcionada por tu administrador', style='List Number')
    doc.add_paragraph('La aplicación cargará la pantalla principal con el listado de tareas', style='List Number')
    
    agregar_placeholder(doc, '📷 PLACEHOLDER: Captura de pantalla de la página principal de la aplicación\n(Mostrar: barra de navegación superior, lista de tareas, botones de acción principales)')
    
    doc.add_heading('Navegación básica', level=2)
    doc.add_paragraph('La aplicación cuenta con cuatro secciones principales, accesibles desde el menú superior:')
    
    secciones = [
        ('Tareas', 'gestión completa de tus tareas pendientes y completadas'),
        ('Categorías', 'crea y gestiona las categorías para organizar tus tareas'),
        ('Plantillas', 'define plantillas reutilizables para crear tareas rápidamente'),
        ('Usuarios', 'administra la lista de usuarios que pueden ser asignados a tareas'),
    ]
    for nombre, desc in secciones:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(nombre).bold = True
        p.add_run(f': {desc}')
    
    agregar_placeholder(doc, '📷 PLACEHOLDER: Captura de pantalla del menú de navegación\n(Incluir: menú superior con las 4 secciones resaltadas)')
    
    doc.add_page_break()

def crear_seccion_tareas(doc):
    """Crea la sección de Gestión de Tareas"""
    doc.add_heading('3. Gestión de Tareas', level=1)
    doc.add_paragraph(
        'Las tareas son el elemento central de la aplicación. Cada tarea tiene un título, un estado '
        '(pendiente o completada), fecha de creación, y puede tener categoría, usuario asignado, '
        'y configuración de repetición.'
    )
    
    doc.add_heading('3.1. Ver todas las tareas', level=2)
    doc.add_paragraph('Haz clic en "Tareas" en el menú superior', style='List Number')
    doc.add_paragraph('Verás un listado con todas las tareas existentes', style='List Number')
    p = doc.add_paragraph('Cada tarea muestra:', style='List Number')
    
    campos = ['Título', 'Estado (pendiente o completada)', 'Categoría asignada (con su color)',
              'Usuario asignado (si lo tiene)', 'Fecha de creación', 'Tipo de recurrencia (si es repetitiva)']
    for campo in campos:
        doc.add_paragraph(campo, style='List Bullet 2')
    
    agregar_placeholder(doc, '📷 PLACEHOLDER: Captura de pantalla de la lista de tareas\n(Mostrar: tabla con varias tareas, algunas completadas, con categorías de colores, usuarios asignados)')
    
    doc.add_heading('3.2. Crear una tarea nueva', level=2)
    doc.add_paragraph('En la página de Tareas, haz clic en el botón "Nueva Tarea"', style='List Number')
    doc.add_paragraph('Se abrirá un formulario con los siguientes campos:', style='List Number')
    
    campos_form = [
        ('Título', 'obligatorio: escribe el nombre de la tarea'),
        ('Categoría', 'opcional: selecciona una categoría del desplegable'),
        ('Usuario asignado', 'opcional: selecciona un usuario responsable'),
        ('¿Es repetitiva?', 'marca esta casilla si quieres que la tarea se repita automáticamente'),
        ('Recurrencia', 'si es repetitiva: selecciona Diaria, Semanal o Mensual'),
    ]
    for nombre, desc in campos_form:
        p = doc.add_paragraph(style='List Bullet 2')
        p.add_run(nombre).bold = True
        p.add_run(f' ({desc})')
    
    doc.add_paragraph('Haz clic en "Guardar"', style='List Number')
    doc.add_paragraph('La tarea aparecerá en el listado principal', style='List Number')
    
    agregar_placeholder(doc, '📷 PLACEHOLDER: Captura de pantalla del formulario de creación de tareas')
    
    doc.add_heading('Ejemplo: Crear una tarea simple', level=3)
    pasos = [
        'Haz clic en "Nueva Tarea"',
        'Introduce el título: "Revisar correo electrónico"',
        'Selecciona categoría: "Trabajo"',
        'Deja sin marcar "Es repetitiva"',
        'Haz clic en "Guardar"'
    ]
    for paso in pasos:
        doc.add_paragraph(paso, style='List Number')
    
    agregar_placeholder(doc, '📷 PLACEHOLDER: Captura del formulario completado con el ejemplo')
    
    doc.add_heading('3.3. Marcar una tarea como completada', level=2)
    doc.add_paragraph('En el listado de tareas, localiza la tarea pendiente', style='List Number')
    doc.add_paragraph('Haz clic en el botón "Completar" (icono de check o marca de verificación)', style='List Number')
    doc.add_paragraph('La tarea se marcará como completada', style='List Number')
    p = doc.add_paragraph(style='List Number')
    p.add_run('Si la tarea es repetitiva').bold = True
    p.add_run(', se creará automáticamente una nueva ocurrencia:')
    
    recurrencias = [
        ('Diaria', 'nueva tarea para mañana'),
        ('Semanal', 'nueva tarea dentro de 7 días'),
        ('Mensual', 'nueva tarea dentro de 30 días'),
    ]
    for tipo, desc in recurrencias:
        p = doc.add_paragraph(style='List Bullet 2')
        p.add_run(tipo).bold = True
        p.add_run(f': {desc}')
    
    agregar_placeholder(doc, '📷 PLACEHOLDER: Captura mostrando tarea antes y después de completar')
    agregar_placeholder(doc, '📊 PLACEHOLDER: Infografía del flujo de tareas repetitivas')
    
    doc.add_page_break()

def crear_seccion_categorias(doc):
    """Crea la sección de Gestión de Categorías"""
    doc.add_heading('4. Gestión de Categorías', level=1)
    doc.add_paragraph(
        'Las categorías te permiten organizar y agrupar tareas por tema, proyecto o prioridad. '
        'Cada categoría tiene un nombre y un color visual que ayuda a identificarlas rápidamente.'
    )
    
    doc.add_heading('4.1. Ver todas las categorías', level=2)
    doc.add_paragraph('Haz clic en "Categorías" en el menú superior', style='List Number')
    doc.add_paragraph('Verás un listado con todas las categorías creadas', style='List Number')
    doc.add_paragraph('Cada categoría muestra: Nombre y Color (representado visualmente)', style='List Number')
    
    agregar_placeholder(doc, '📷 PLACEHOLDER: Captura de pantalla de la lista de categorías')
    
    doc.add_heading('4.2. Crear una categoría nueva', level=2)
    doc.add_paragraph('En la página de Categorías, haz clic en "Nueva Categoría"', style='List Number')
    doc.add_paragraph('Se abrirá un formulario con los siguientes campos:', style='List Number')
    
    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run('Nombre').bold = True
    p.add_run(' (obligatorio): escribe el nombre de la categoría')
    
    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run('Color').bold = True
    p.add_run(' (obligatorio): selecciona un color usando el selector visual o introduce un código hexadecimal (ejemplo: #FF5733)')
    
    doc.add_paragraph('Haz clic en "Guardar"', style='List Number')
    doc.add_paragraph('La categoría aparecerá en el listado y estará disponible para asignar a tareas', style='List Number')
    
    agregar_placeholder(doc, '📷 PLACEHOLDER: Captura de pantalla del formulario de creación de categorías')
    
    doc.add_page_break()

def crear_seccion_plantillas(doc):
    """Crea la sección de Gestión de Plantillas"""
    doc.add_heading('5. Gestión de Plantillas', level=1)
    doc.add_paragraph(
        'Las plantillas son configuraciones reutilizables que te permiten crear tareas rápidamente '
        'con valores predefinidos. Son especialmente útiles para tareas repetitivas que creas con frecuencia.'
    )
    
    doc.add_heading('5.1. Ver todas las plantillas', level=2)
    doc.add_paragraph('Haz clic en "Plantillas" en el menú superior', style='List Number')
    doc.add_paragraph('Verás un listado con todas las plantillas creadas', style='List Number')
    
    agregar_placeholder(doc, '📷 PLACEHOLDER: Captura de pantalla de la lista de plantillas')
    
    doc.add_heading('5.2. Crear una plantilla nueva', level=2)
    doc.add_paragraph('En la página de Plantillas, haz clic en "Nueva Plantilla"', style='List Number')
    doc.add_paragraph('Se abrirá un formulario', style='List Number')
    doc.add_paragraph('Haz clic en "Guardar"', style='List Number')
    
    agregar_placeholder(doc, '📷 PLACEHOLDER: Captura del formulario de creación de plantillas')
    
    doc.add_heading('5.3. Crear una tarea desde una plantilla (instanciar)', level=2)
    doc.add_paragraph('Esta es la función principal de las plantillas: crear tareas rápidamente con un solo clic.')
    doc.add_paragraph('En el listado de plantillas, localiza la plantilla que quieres usar', style='List Number')
    doc.add_paragraph('Haz clic en el botón "Instanciar" o "Crear Tarea"', style='List Number')
    doc.add_paragraph('Se creará automáticamente una nueva tarea', style='List Number')
    
    agregar_placeholder(doc, '📊 PLACEHOLDER: Infografía del flujo de instanciación de plantillas')
    
    doc.add_page_break()

def crear_seccion_usuarios(doc):
    """Crea la sección de Gestión de Usuarios"""
    doc.add_heading('6. Gestión de Usuarios Asignados', level=1)
    doc.add_paragraph(
        'Los usuarios asignados son las personas que pueden ser responsables de las tareas. '
        'Puedes crear, editar y eliminar usuarios, y luego asignarlos a tareas específicas.'
    )
    
    doc.add_heading('6.1. Ver todos los usuarios', level=2)
    doc.add_paragraph('Haz clic en "Usuarios" en el menú superior', style='List Number')
    doc.add_paragraph('Verás un listado con todos los usuarios registrados', style='List Number')
    
    agregar_placeholder(doc, '📷 PLACEHOLDER: Captura de pantalla de la lista de usuarios')
    
    doc.add_heading('6.2. Crear un usuario nuevo', level=2)
    doc.add_paragraph('En la página de Usuarios, haz clic en "Nuevo Usuario"', style='List Number')
    doc.add_paragraph('Se abrirá un formulario', style='List Number')
    doc.add_paragraph('Haz clic en "Guardar"', style='List Number')
    
    agregar_placeholder(doc, '📷 PLACEHOLDER: Captura del formulario de creación de usuarios')
    
    doc.add_page_break()

def crear_seccion_faq(doc):
    """Crea la sección de Preguntas Frecuentes"""
    doc.add_heading('7. Preguntas frecuentes', level=1)
    
    faqs = [
        ('¿Qué pasa si completo una tarea repetitiva?',
         'Al marcar como completada una tarea repetitiva, se crean dos cosas: '
         '(1) La tarea actual queda marcada como completada (historial). '
         '(2) Se crea automáticamente una nueva tarea pendiente con la fecha calculada.'),
        
        ('¿Puedo eliminar una categoría que tiene tareas?',
         'No. Para eliminar una categoría, primero debes reasignar las tareas a otra categoría, '
         'eliminar las tareas que usan esa categoría, o editar las tareas para quitarles la categoría.'),
        
        ('¿Qué diferencia hay entre una plantilla y una tarea repetitiva?',
         'Plantilla: es una configuración guardada que usas manualmente. '
         'Tarea repetitiva: se genera automáticamente al completarse sin intervención manual.'),
        
        ('¿Se puede cambiar una tarea normal a repetitiva?',
         'Sí. Edita la tarea, marca la casilla "Es repetitiva", selecciona la recurrencia, y guarda.'),
        
        ('¿Qué pasa si elimino un usuario que tiene tareas asignadas?',
         'El usuario se elimina, pero las tareas no. Las tareas quedan sin usuario asignado.'),
    ]
    
    for pregunta, respuesta in faqs:
        p = doc.add_paragraph()
        p.add_run(pregunta).bold = True
        doc.add_paragraph(respuesta)
        doc.add_paragraph()
    
    doc.add_page_break()

def crear_seccion_glosario(doc):
    """Crea la sección de Glosario"""
    doc.add_heading('8. Glosario', level=1)
    
    terminos = [
        ('Tarea', 'elemento individual de trabajo pendiente o completado'),
        ('Categoría', 'etiqueta temática con color para agrupar tareas'),
        ('Plantilla', 'configuración reutilizable para crear tareas rápidamente'),
        ('Usuario asignado', 'persona responsable de una tarea'),
        ('Tarea repetitiva', 'tarea que se regenera automáticamente al completarse'),
        ('Recurrencia', 'periodicidad de repetición (diaria, semanal, mensual)'),
        ('Instanciar', 'crear una nueva tarea a partir de una plantilla'),
    ]
    
    for termino, definicion in terminos:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(termino).bold = True
        p.add_run(f': {definicion}')
    
    doc.add_page_break()

def crear_seccion_soporte(doc):
    """Crea la sección de Soporte"""
    doc.add_heading('9. Soporte y contacto', level=1)
    doc.add_paragraph(
        'Si tienes problemas técnicos, dudas sobre el uso de la aplicación, '
        'o sugerencias de mejora, contacta con:'
    )
    
    contactos = [
        ('Soporte técnico', 'soporte@democopilot.com'),
        ('Equipo de desarrollo', 'desarrollo@democopilot.com'),
        ('Documentación online', 'https://democopilot.com/docs'),
    ]
    
    for tipo, contacto in contactos:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(tipo).bold = True
        p.add_run(f': {contacto}')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Fin del manual
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Fin del manual de usuario')
    run.italic = True
    run.font.color.rgb = RGBColor(102, 102, 102)

def main():
    """Función principal que genera el documento"""
    print("🔄 Generando manual de usuario en formato Word...")
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos básicos
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    
    # Generar contenido
    crear_portada(doc)
    
    # Tabla de contenidos (se debe actualizar manualmente en Word)
    doc.add_paragraph('Tabla de Contenidos')
    doc.add_paragraph('(Actualizar esta tabla en Word: clic derecho → Actualizar campos)')
    doc.add_page_break()
    
    crear_seccion_intro(doc)
    crear_seccion_primeros_pasos(doc)
    crear_seccion_tareas(doc)
    crear_seccion_categorias(doc)
    crear_seccion_plantillas(doc)
    crear_seccion_usuarios(doc)
    crear_seccion_faq(doc)
    crear_seccion_glosario(doc)
    crear_seccion_soporte(doc)
    
    # Guardar documento
    doc.save('manual-usuario.docx')
    
    print("✅ Manual de usuario generado: docs/manual-usuario.docx")
    print()
    print("📊 Estadísticas:")
    print("  - Formato: Word (.docx)")
    print("  - Tipo: Documentación completa")
    print("  - Secciones: 9")
    print("  - Funcionalidades documentadas: 15+")
    print("  - Placeholders para capturas: 18")
    print()
    print("📝 Próximos pasos:")
    print("  1. Abrir el documento en Word")
    print("  2. Actualizar la tabla de contenidos (clic derecho → Actualizar campos)")
    print("  3. Reemplazar los placeholders con capturas de pantalla reales")
    print("  4. Validar los pasos descritos ejecutándolos en la aplicación")
    print("  5. Compartir con el equipo para feedback")

if __name__ == '__main__':
    main()
